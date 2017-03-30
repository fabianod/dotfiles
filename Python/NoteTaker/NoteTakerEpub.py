import textract
import random

from docx import Document;

text = textract.process('To The Light House.epub')
documentName = "To The Light House.docx"

text = text.split("\n")

counter = 1
chapters = []
chapter = ""
gotToChapter1 = False

for line in text:
    chapterCounter = "Chapter " + str(counter)
    if line == chapterCounter: 
        gotToChapter1 = True
        counter += 1
        if gotToChapter1:
            chapters.append(chapter)
            chapter = ""
    if len(line) > 0:
        chapter += line

counter = 1
document = Document()

failCounter = 0;
def randomNote(lines):

    global failCounter

    if failCounter > 5:
        failCounter = 0
        return None;
    
    randomNoteLine = random.choice(lines)

    if len(randomNoteLine.split(" ")) <= 5:
        failCounter += 1
        return randomNote(lines)
    else:
        failCounter = 0
        return randomNoteLine + "."

def slicePer(source, step):
    return [source[i::step] for i in range(step)]

for chapter in chapters:

    chapterLine = document.add_paragraph('', style='ListBullet')
    chapterLine.add_run("Chapter " + str(counter)).bold=True
    lines = chapter.split(".")

    linesSplit = slicePer(lines, 10)

    for lineSplit in linesSplit:
        try:
            randomLine = randomNote(lineSplit)
            if randomLine is None:
                break;

            randomLine = randomLine.encode('ascii', 'ignore')
            noteLine = document.add_paragraph('', style='ListBullet')
            print(randomLine)
            noteLine.add_run(randomLine).bold = False;
        except (Exception):
            print("ERROR: " + randomLine)

    counter += 1

document.save(documentName)
print("Notes created")

